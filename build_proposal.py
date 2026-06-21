# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x34, 0x78, 0xF6)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x4A, 0x55, 0x68)
MUTED = RGBColor(0x7A, 0x8A, 0x9E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHTBLUE = "E8F0FE"
BLUEHEX = "3478F6"
WARNBG = "FFF4E5"

doc = Document()

# Base font
style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
style.font.size = Pt(10.5)
style.font.color.rgb = DARK

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_cell_font(cell, size=10, bold=False, color=DARK, align=None):
    for p in cell.paragraphs:
        if align:
            p.alignment = align
        for r in p.runs:
            r.font.name = 'Malgun Gothic'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
            r.font.size = Pt(size)
            r.bold = bold
            r.font.color.rgb = color

def add_heading(text, size=15, color=BLUE, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Malgun Gothic'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    r.font.size = Pt(size)
    r.bold = True
    r.font.color.rgb = color
    return p

def add_body(text, bold=False, color=DARK, size=10.5, space_after=4, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    r = p.add_run(text)
    r.font.name = 'Malgun Gothic'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    return p

def add_bullet(text, color=GRAY, size=10):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• " + text)
    r.font.name = 'Malgun Gothic'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def style_table(table):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def header_row(table, headers):
    cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cells[i].text = h
        shade(cells[i], BLUEHEX)
        set_cell_font(cells[i], size=10, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

# ===================== COVER =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(30)
r = title.add_run("HighPlus")
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
r.font.size = Pt(13)
r.bold = True
r.font.color.rgb = BLUE

h = doc.add_paragraph()
h.paragraph_format.space_after = Pt(2)
r = h.add_run("거래처 마케팅 진단 · 실행 제안서")
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
r.font.size = Pt(24)
r.bold = True
r.font.color.rgb = DARK

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("참나무요양병원  ·  까치한의원(날씬하냑)")
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
r.font.size = Pt(13)
r.font.color.rgb = GRAY

d = doc.add_paragraph()
r = d.add_run("전략실 논의 정리 (6/17 · 6/19)  |  작성: HighPlus Marketing")
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

# divider
p = doc.add_paragraph()
pPr = p._p.get_or_add_pPr()
pbdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12'); bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), BLUEHEX)
pbdr.append(bottom); pPr.append(pbdr)

# ===================== 0. 개요 =====================
add_heading("0. 제안 개요", size=16)
add_body("두 거래처의 채널 현황을 진단하고, 가장 효과가 빠른 순서대로 실행안을 제안합니다. "
         "참나무요양병원은 '플레이스 재세팅', 까치한의원은 '계정 통합 + 주력 키워드 1개 확정'이 각 건의 첫 단추입니다.",
         color=GRAY, space_after=8)

t = doc.add_table(rows=1, cols=3)
style_table(t)
header_row(t, ["구분", "참나무요양병원", "까치한의원 (날씬하냑)"])
rows = [
    ("업종", "요양병원", "한의원 (다이어트 한약)"),
    ("핵심 채널", "플레이스 · 블로그 · 인스타", "인스타 · 블로그 · 체험단"),
    ("가장 큰 문제", "플레이스 키워드 방치 · 인스타 중단", "계정 정체성 혼선 · 키워드 분산"),
    ("1순위 액션", "플레이스 재세팅", "인스타 계정 통합 + 주력 키워드 1개 확정"),
    ("리스크", "요양병원 블로그 수익성(돈 날릴 위험)", "의료광고법(치료경험담) 위반 소지"),
]
for label, a, b in rows:
    cells = t.add_row().cells
    cells[0].text = label; shade(cells[0], LIGHTBLUE); set_cell_font(cells[0], size=9.5, bold=True, color=DARK)
    cells[1].text = a; set_cell_font(cells[1], size=9.5)
    cells[2].text = b; set_cell_font(cells[2], size=9.5)

# ===================== 1. 참나무요양병원 =====================
doc.add_page_break()
add_heading("1. 참나무요양병원", size=16)
add_body("플레이스 6위권(1페이지 진입 임박), 자체 블로그·홈페이지는 운영 중이나 자체 인스타는 멈춘 상태. "
         "신규로 맡고 싶은 건이지만 요양병원 특성상 수익성 검토가 필요합니다.", color=GRAY, space_after=8)

add_heading("채널 현황 진단", size=12, space_before=8)
t = doc.add_table(rows=1, cols=3)
style_table(t)
header_row(t, ["채널", "현황", "진단 / 개선 포인트"])
diag = [
    ("플레이스", "오늘 기준 6위권", "1페이지 진입 임박. 대표키워드 5개 중 3개를 버리는 중 → 재세팅 필요. 대표사진이 전부 건물 → 사람 나오는 사진으로 차별화"),
    ("블로그", "자체 운영 중 (월 약 30개, 하루 1개)", "원내 프로그램 위주(인지바달, 사회복지 게임·예배·미용 만들기 등). 스킨 버튼(병원소개/재활요양/암면역) 눌러도 이미지 1장+4줄뿐 → 환자 궁금증·의심 해소 안 됨"),
    ("인스타", "운영 중단", "사진·식단·프로그램 모으는 채널로 중요. 재가동 필요(직접 어려우면 우리가 사진 받아 대행)"),
    ("홈페이지", "운영 중", "살짝 올드하지만 무난"),
]
for ch, st, dg in diag:
    cells = t.add_row().cells
    cells[0].text = ch; shade(cells[0], LIGHTBLUE); set_cell_font(cells[0], size=9.5, bold=True, color=DARK)
    cells[1].text = st; set_cell_font(cells[1], size=9.5)
    cells[2].text = dg; set_cell_font(cells[2], size=9.5)

add_heading("핵심 이슈 — 수익성", size=12, space_before=10)
add_bullet("요양병원은 블로그 인기글이 지수대로 노출되지 않아, 잘못하면 돈만 날리는 구조.")
add_bullet("경쟁사 금솔: 월 300 → 100으로 축소, 월 기사 1건만 발행하기로 함. → 우리 단가·범위 사전 확정 필요.")

add_heading("실행 우선순위", size=12, space_before=10)
steps = [
    ("1. 플레이스 재세팅", "가장 빠른 효과. 버리고 있는 키워드 3개 살리기 + 정보 보충 + 사람 나오는 대표사진으로 교체."),
    ("2. 전용 블로그 1개 신설", "질환/치료/시설 설명 중심. 환자 궁금증·의심을 풀어주는 콘텐츠로 운영."),
    ("3. 인스타 재가동", "사진·식단·프로그램 업로드 채널로 살리기. 필요시 우리가 사진 받아 대행."),
    ("4. 단가·범위 사전 확정", "요양병원 수익성 리스크 검토 후 진입. 경쟁사 사례 참고해 단가 결정."),
]
for head, desc in steps:
    add_body(head, bold=True, color=BLUE, size=11, space_after=1)
    add_body(desc, color=GRAY, size=10, space_after=6, indent=10)

# ===================== 2. 까치한의원 =====================
doc.add_page_break()
add_heading("2. 까치한의원 (날씬하냑 / 하냑최선생)", size=16)
add_body("인스타·블로그 모두 노출이 안 나오는 상태. 가장 큰 문제는 계정 정체성 혼선과 키워드 분산입니다.",
         color=GRAY, space_after=8)

add_heading("현황 진단", size=12, space_before=8)
t = doc.add_table(rows=1, cols=2)
style_table(t)
header_row(t, ["항목", "현황 / 진단"])
diag2 = [
    ("인스타", "체험단 배포만 했을 뿐 알맹이 없음. 카드뉴스 톤앤매너 엉망, 피드 발행 들쭉날쭉. '하냑최선생' 계정에 원장님까지 나와 정체성 혼선 → 알고리즘이 타겟을 못 잡아 노출 안 됨. 계정 2개라 혼선 → 하나 정리 필요"),
    ("블로그", "지수 등은 본부에서 파악 후 공유 (대기 중)"),
    ("키워드 분산", "'하냑최선생 / 날씬하냑 / 날씬한정 / 날씬한팩' 4개로 흩어짐"),
]
for k, v in diag2:
    cells = t.add_row().cells
    cells[0].text = k; shade(cells[0], LIGHTBLUE); set_cell_font(cells[0], size=9.5, bold=True, color=DARK)
    cells[1].text = v; set_cell_font(cells[1], size=9.5)

add_heading("6/19 날씬하냑 추가 분석 (김선수)", size=12, space_before=10, color=BLUE)
add_bullet("체험단 효과 분석 — 블로그 체험단은 주로 준최 4~5짜리로 배포. 5월 약 20건, 4월 약 3건. 전국 살포(홍대·신림·창원·울산·부산·원주 등 'OO다이어트한약/한의원')했지만 매출 계속 하락 → 효과 없는 방식으로 판단.")
add_bullet("핵심 발견 — '날씬하냑' 검색 시 '날씬한약'으로 자동 정정됨 → 브랜드명 검색하는 사람이 없음(월 30회 미만). 브랜드 인지도 바닥.")
add_bullet("방향 전환 — 체험단 살포 → 자체 네이버 블로그 운영으로 브랜드 강화. 체험단은 하더라도 소량(월 10회 이하).")
add_bullet("필수 2가지 — ① 다이어트 콘텐츠 꾸준히 쓸 사람 확보 ② 해당 지역 블로그 섹터가 뜨는지 확인.")
add_bullet("블로그 제목 전략(김태헌) — '날씬하냑(까치한의원 다이어트)'처럼 제목 두 개 병기. 초진·매출 증가에 초점.")

add_heading("실행 우선순위", size=12, space_before=10)
steps2 = [
    ("1. 인스타 계정 1개로 통합 + 컨셉 고정", "한약/원장님/하냑최선생이 섞이지 않게 계정 정체성을 하나로 명확히."),
    ("2. 자체 네이버 블로그 운영으로 전환", "체험단 의존 탈피. 다이어트 검색자에게 조금씩 인지시키는 방식."),
    ("3. 키워드 교통정리", "날씬하냑을 상위 우산 브랜드로, 제품 중 주력 1개를 정해 그 키워드부터 집중 공략."),
    ("4. 블로그 제목 병기", "'날씬하냑 + 까치한의원 다이어트'로 검색 유입 + 브랜드 동시 노출."),
    ("5. 입점·바이럴 검토", "모두닥/강남언니 등은 의료광고법 체크 후 진행."),
]
for head, desc in steps2:
    add_body(head, bold=True, color=BLUE, size=11, space_after=1)
    add_body(desc, color=GRAY, size=10, space_after=6, indent=10)

# Warning box
add_heading("⚠ 의료광고법 유의", size=11, space_before=8, color=RGBColor(0xC0,0x6A,0x00))
wt = doc.add_table(rows=1, cols=1)
wt.style = 'Table Grid'
wcell = wt.rows[0].cells[0]
shade(wcell, WARNBG)
wcell.text = ("한약 후기 바이럴·인스타 체험단은 의료광고법(치료경험담) 위반 소지가 있습니다. "
              "카페바이럴 침투(30~40건)·어플 연계 집행 전 반드시 법적 체크 필요.")
set_cell_font(wcell, size=9.5, color=RGBColor(0x8A,0x5A,0x00))

# 입점 참고표
add_heading("입점·바이럴 참고 (강남언니 기준)", size=11, space_before=10)
t = doc.add_table(rows=1, cols=2)
style_table(t)
header_row(t, ["항목", "비용"])
fees = [("입점비", "220만원"), ("프리미엄", "월 100만원"), ("메인배너", "월 300만원"),
        ("타겟", "20~50대 (예산 여유 시)"), ("바이럴", "카페·바이럴 30~40건 + 어플 연계")]
for k, v in fees:
    cells = t.add_row().cells
    cells[0].text = k; shade(cells[0], LIGHTBLUE); set_cell_font(cells[0], size=9.5, bold=True, color=DARK)
    cells[1].text = v; set_cell_font(cells[1], size=9.5)

# ===================== 3. 종합 =====================
doc.add_page_break()
add_heading("3. 종합 — 거래처별 1순위 액션", size=16)
t = doc.add_table(rows=1, cols=2)
style_table(t)
header_row(t, ["거래처", "지금 당장 할 첫 단추"])
summ = [
    ("참나무요양병원", "플레이스 재세팅 (버린 키워드 3개 살리기 + 사람 나오는 대표사진)"),
    ("까치한의원(날씬하냑)", "인스타 계정 통합 + 주력 키워드 1개 확정 + 자체 블로그 전환"),
]
for k, v in summ:
    cells = t.add_row().cells
    cells[0].text = k; shade(cells[0], LIGHTBLUE); set_cell_font(cells[0], size=10, bold=True, color=BLUE)
    cells[1].text = v; set_cell_font(cells[1], size=10)

doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("HighPlus Marketing  ·  highplus.work")
r.font.name = 'Malgun Gothic'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
r.font.size = Pt(9)
r.font.color.rgb = MUTED

out = "/home/user/highplus-marketing/outputs/거래처_마케팅_제안서.docx"
doc.save(out)
print("saved:", out)
